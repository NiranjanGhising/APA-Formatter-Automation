"""
APA Formatter Core Service
Handles the main formatting logic for APA 7th edition

APA 7th Edition Guidelines (from Scribbr, Purdue OWL & APA Manual):
- 1 inch margins on all sides
- Double-spaced all text including headings
- First line indent 0.5 inches for body paragraphs (skip if already indented)
- Font: Times New Roman 12pt, Arial 11pt, or Georgia 11pt
- ALL text must be BLACK color
- Page numbers on every page (right-aligned header)

Heading Levels (APA 7th):
- Level 1: Centered, Bold, Title Case - START ON NEW PAGE
- Level 2: Flush Left, Bold, Title Case  
- Level 3: Flush Left, Bold Italic, Title Case
- Level 4: Indented 0.5", Bold, Title Case, ends with period
- Level 5: Indented 0.5", Bold Italic, Title Case, ends with period

Front Matter (student papers):
- Table of Contents: "Contents" centered, NORMAL font (not bold)
- List of Tables: "List of Tables" NORMAL font (not bold)
- List of Figures: "List of Figures" NORMAL font (not bold)
- All entries include page numbers with dot leaders

Figure/Table Captions:
- "Figure X" or "Table X" = BOLD only (no italic)
- Title = ITALIC only (no bold)

IMPORTANT: 
- Do NOT duplicate existing cover pages, TOC, LOF, LOT
- Do NOT add indent if paragraph already has 0.5" indent
- New major sections (Level 1) start on new page
"""

from docx import Document
from docx.shared import Pt, Inches, Twips, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from typing import List, Dict, Optional, Tuple
import re
import copy
from lxml import etree
import docx


class APAFormatter:
    """Core APA 7th Edition formatting engine - preserves document content"""
    
    # APA 7th Edition constants
    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(12)
    MARGIN = Inches(1)
    LINE_SPACING = 2.0  # Double-spaced
    FIRST_LINE_INDENT = Inches(0.5)
    HANGING_INDENT = Inches(0.5)
    BLACK_COLOR = RGBColor(0, 0, 0)
    
    # Indent tolerance for checking existing indents (0.4" to 0.6")
    INDENT_TOLERANCE_MIN = Inches(0.4)
    INDENT_TOLERANCE_MAX = Inches(0.6)
    
    # Common heading keywords for detection
    LEVEL_1_KEYWORDS = [
        'abstract', 'introduction', 'method', 'methods', 'methodology',
        'results', 'discussion', 'conclusion', 'conclusions', 'references',
        'bibliography', 'appendix', 'acknowledgements', 'acknowledgments',
        'literature review', 'theoretical framework', 'findings',
        'recommendations', 'limitations', 'future research', 'summary',
        'table of contents', 'list of tables', 'list of figures',
        'chapter', 'background', 'problem statement', 'research questions',
        'hypothesis', 'hypotheses', 'definition of terms', 'scope',
        'significance', 'purpose', 'objectives', 'overview', 'contents'
    ]
    
    LEVEL_2_KEYWORDS = [
        'participants', 'materials', 'procedure', 'measures', 'design',
        'data analysis', 'data collection', 'sampling', 'instruments',
        'variables', 'ethical considerations', 'reliability', 'validity',
        'demographics', 'descriptive statistics', 'inferential statistics',
        'qualitative analysis', 'quantitative analysis', 'themes',
        'subthemes', 'categories', 'implications', 'contributions',
        'previous studies', 'current study', 'research design'
    ]
    
    def __init__(self, source_path: str):
        self.source_path = source_path
        self.doc = Document(source_path)
        self.issues_found = []
        self.issues_fixed = 0
        # Storage for TOC, LOF, LOT generation
        self.headings_found = []  # [(level, text, para_index)]
        self.figures_found = []   # [(num, title)]
        self.tables_found = []    # [(num, title)]
        # Track existing elements
        self.has_title_page = False
        self.has_toc = False
        self.has_lof = False
        self.has_lot = False
    
    def format_document(self, options: Dict) -> Document:
        """Main entry point - format the document while preserving content"""
        
        # Step 0: Detect existing elements FIRST (before any modifications)
        self._detect_existing_elements()
        
        # Step 1: Set up document margins and default styles
        self._setup_document_styles()
        
        # Step 2: Add page numbers to document
        self._add_page_numbers()
        
        # Step 3: First pass - detect and mark all headings
        self._detect_all_headings()
        
        # Step 4: Format headings with page breaks for Level 1
        if options.get('fix_headings', True):
            self._format_all_headings_with_page_breaks()
        
        # Step 5: Ensure all fonts are black
        self._ensure_all_fonts_black()
        
        # Step 6: Format body paragraphs (with indent check)
        if options.get('fix_spacing', True):
            self._format_body_paragraphs()
        
        # Step 7: Clean up excessive blank paragraphs
        self._remove_excessive_blanks()
        
        # Step 8: Format references section if present
        if options.get('fix_references', True):
            self._format_references_section()
        
        # Step 9: Detect and format tables and figures
        self._detect_tables_and_figures()
        self._format_tables_and_figures()
        
        # Step 10: Generate front matter (TOC, LOF, LOT) if requested AND not existing
        if options.get('generate_toc', True):
            self._generate_front_matter()
        
        return self.doc
    
    # ==================== DETECTION OF EXISTING ELEMENTS ====================
    
    def _detect_existing_elements(self):
        """Detect if document already has title page, TOC, LOF, LOT"""
        first_page_content = []
        found_page_break = False
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip().lower()
            
            # Collect first page content (before first page break or first 10 paras)
            if not found_page_break and i < 15:
                # Check for page break
                if self._has_page_break(para):
                    found_page_break = True
                first_page_content.append(text)
            
            # Check for TOC
            if text in ['contents', 'table of contents']:
                self.has_toc = True
            
            # Check for LOF
            if text == 'list of figures':
                self.has_lof = True
            
            # Check for LOT
            if text == 'list of tables':
                self.has_lot = True
        
        # Detect title page - look for typical title page patterns
        self.has_title_page = self._detect_title_page(first_page_content)
    
    def _detect_title_page(self, first_page_content: List[str]) -> bool:
        """Detect if document has a title page based on content patterns"""
        if not first_page_content:
            return False
        
        # Common title page indicators
        title_page_patterns = [
            r'(university|college|institute|school)',
            r'(department|faculty|program)',
            r'(submitted|presented|prepared)',
            r'(course|class|assignment)',
            r'(professor|instructor|supervisor)',
            r'\d{4}',  # Year
            r'(january|february|march|april|may|june|july|august|september|october|november|december)',
        ]
        
        content_joined = ' '.join(first_page_content[:10])
        
        # Check how many patterns match
        matches = sum(1 for pattern in title_page_patterns 
                     if re.search(pattern, content_joined, re.IGNORECASE))
        
        # If 3+ patterns match, likely a title page
        return matches >= 3
    
    def _has_page_break(self, para) -> bool:
        """Check if paragraph contains a page break"""
        for run in para.runs:
            if run._r.xml and 'w:br' in run._r.xml and 'w:type="page"' in run._r.xml:
                return True
        # Also check paragraph XML directly
        p_xml = para._p.xml
        return 'w:br w:type="page"' in p_xml or '<w:br w:type="page"/>' in p_xml
    
    def _setup_document_styles(self):
        """Set up document with APA formatting - margins and default font"""
        # Set margins for all sections
        for section in self.doc.sections:
            section.top_margin = self.MARGIN
            section.bottom_margin = self.MARGIN
            section.left_margin = self.MARGIN
            section.right_margin = self.MARGIN
        
        # Update the Normal style
        try:
            style = self.doc.styles['Normal']
            font = style.font
            font.name = self.FONT_NAME
            font.size = self.FONT_SIZE
            font.color.rgb = self.BLACK_COLOR
            
            pf = style.paragraph_format
            pf.line_spacing = self.LINE_SPACING
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)
        except:
            pass
        
        self.issues_fixed += 1
    
    def _add_page_numbers(self):
        """Add page numbers to the document header (right-aligned)"""
        for section in self.doc.sections:
            header = section.header
            # Clear existing header content but keep structure
            if not header.paragraphs:
                header.add_paragraph()
            
            # Use first paragraph in header
            header_para = header.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add page number field
            run = header_para.add_run()
            self._add_page_number_field(run)
            
            # Apply font
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE
            run.font.color.rgb = self.BLACK_COLOR
        
        self.issues_fixed += 1
    
    def _add_page_number_field(self, run):
        """Add a PAGE field to a run for automatic page numbering"""
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_separate)
        run._r.append(fldChar_end)
    
    def _ensure_all_fonts_black(self):
        """Ensure all text in the document is black color"""
        # Process all paragraphs
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.color.rgb != self.BLACK_COLOR:
                    run.font.color.rgb = self.BLACK_COLOR
        
        # Process all tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.color.rgb != self.BLACK_COLOR:
                                run.font.color.rgb = self.BLACK_COLOR
        
        self.issues_fixed += 1
    
    def _detect_all_headings(self):
        """First pass: detect all headings and store them"""
        self.headings_found = []
        
        for idx, para in enumerate(self.doc.paragraphs):
            level = self._detect_heading_level(para)
            if level:
                text = para.text.strip()
                self.headings_found.append((level, text, idx))
    
    def _format_all_headings_with_page_breaks(self):
        """Format all headings with APA 5-level system, adding page breaks for Level 1"""
        # Track if we've seen the first Level 1 heading (don't add page break before first one)
        first_level1_seen = False
        
        for level, text, idx in self.headings_found:
            para = self.doc.paragraphs[idx]
            
            # For Level 1 headings (except the very first one), add page break before
            if level == 1:
                text_lower = text.lower().strip()
                # Skip page break for front matter items
                if text_lower not in ['contents', 'table of contents', 'list of figures', 'list of tables']:
                    if first_level1_seen:
                        self._insert_page_break_before(para)
                    first_level1_seen = True
            
            self._apply_heading_format(para, level)
            self.issues_fixed += 1
    
    def _insert_page_break_before(self, para):
        """Insert a page break before a paragraph"""
        # Create a new paragraph with page break and insert before target
        p_element = para._p
        
        # Create page break paragraph
        page_break_p = OxmlElement('w:p')
        run = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run.append(br)
        page_break_p.append(run)
        
        # Insert before the heading paragraph
        p_element.addprevious(page_break_p)
    
    def _detect_heading_level(self, para) -> Optional[int]:
        """
        Detect if paragraph is a heading and what level using multiple heuristics:
        1. Built-in Word heading styles
        2. Font size differences (larger = more important)
        3. Formatting (bold, centered, all caps)
        4. Content matching (common heading keywords)
        5. Length (headings are typically short)
        6. Position context (after page breaks, etc.)
        """
        text = para.text.strip()
        if not text:
            return None
        
        # Skip very long text - not a heading
        if len(text) > 150:
            return None
        
        # Skip if it looks like a caption
        text_lower = text.lower()
        if text_lower.startswith(('figure ', 'table ', 'note.', 'note:')):
            return None
        
        # ============ CHECK 1: Built-in heading styles ============
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            # Check for "Heading 1", "Heading 2", etc.
            for i in range(1, 10):
                if f'heading {i}' in style_name or f'heading{i}' in style_name:
                    return min(i, 5)
            # Also check for "Title" style
            if style_name == 'title':
                return 1
            # Check for "Subtitle"
            if style_name == 'subtitle':
                return 2
        
        # ============ CHECK 2: Gather formatting info ============
        is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        is_bold = False
        is_italic = False
        is_underlined = False
        font_size = None
        
        if para.runs:
            non_empty_runs = [r for r in para.runs if r.text.strip()]
            if non_empty_runs:
                is_bold = all(r.bold for r in non_empty_runs)
                is_italic = all(r.italic for r in non_empty_runs)
                is_underlined = all(r.underline for r in non_empty_runs)
                # Get font size from first run
                if non_empty_runs[0].font.size:
                    font_size = non_empty_runs[0].font.size.pt
        
        is_all_caps = text.isupper() and len(text) > 2 and len(text) < 100
        is_short = len(text) < 80
        
        # ============ CHECK 3: Content-based detection ============
        text_clean = text.lower().strip().rstrip('.:')
        
        # Check if it matches Level 1 keywords exactly or starts with them
        is_level1_keyword = False
        for keyword in self.LEVEL_1_KEYWORDS:
            if text_clean == keyword or text_clean.startswith(keyword + ' '):
                is_level1_keyword = True
                break
            # Also check for numbered chapters: "Chapter 1", "1. Introduction"
            if re.match(r'^(chapter\s*\d+|^\d+\.?\s*' + re.escape(keyword) + r')$', text_clean):
                is_level1_keyword = True
                break
        
        # Check Level 2 keywords
        is_level2_keyword = False
        for keyword in self.LEVEL_2_KEYWORDS:
            if text_clean == keyword or text_clean.startswith(keyword + ' '):
                is_level2_keyword = True
                break
        
        # ============ CHECK 4: Font size heuristic ============
        # If font is significantly larger than 12pt, likely a heading
        is_large_font = font_size and font_size > 13
        
        # ============ DECISION LOGIC ============
        
        # Strong indicators of Level 1
        if is_level1_keyword and (is_centered or is_bold or is_all_caps or is_large_font):
            return 1
        
        # Centered + Bold + Short = Level 1
        if is_centered and is_bold and is_short:
            return 1
        
        # All caps + short = Level 1
        if is_all_caps and is_short:
            return 1
        
        # Large font + short = Level 1
        if is_large_font and is_short and (is_bold or is_centered):
            return 1
        
        # Level 1 keyword alone (even without special formatting in messy docs)
        if is_level1_keyword and is_short:
            return 1
        
        # Level 2 indicators - keyword with formatting OR short keyword alone
        if is_level2_keyword and (is_bold or is_large_font):
            return 2
        
        # Level 2 keyword alone (for messy docs without formatting) if SHORT
        if is_level2_keyword and is_short and len(text) < 50:
            return 2
        
        # Bold + Italic + Left-aligned = Level 3
        if is_bold and is_italic and not is_centered and is_short:
            return 3
        
        # Bold only + Left-aligned + Short = Level 2
        if is_bold and not is_centered and not is_italic and is_short:
            return 2
        
        # Underlined headings (common in messy docs) - treat as Level 2
        if is_underlined and is_short and not is_bold:
            return 2
        
        # Numbered sections like "1.1", "2.3.1" etc.
        if re.match(r'^\d+(\.\d+)*\.?\s+\w', text):
            # Count dots to determine level
            match = re.match(r'^(\d+(?:\.\d+)*)', text)
            if match:
                num_parts = len(match.group(1).split('.'))
                return min(num_parts, 5)
        
        return None
    
    def _apply_heading_format(self, para, level: int):
        """Apply proper APA heading format to a paragraph"""
        text = para.text.strip()
        
        # Convert to title case
        formatted_text = self._to_title_case(text)
        
        # Clear existing formatting first
        para.clear()
        
        # Set paragraph formatting based on level
        # IMPORTANT: No extra space before/after headings in APA!
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = self.LINE_SPACING
        
        if level == 1:
            # Level 1: Centered, Bold, Title Case
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Inches(0)
            run = para.add_run(formatted_text)
            run.bold = True
            run.italic = False
            
        elif level == 2:
            # Level 2: Flush Left, Bold, Title Case
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Inches(0)
            run = para.add_run(formatted_text)
            run.bold = True
            run.italic = False
            
        elif level == 3:
            # Level 3: Flush Left, Bold Italic, Title Case
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Inches(0)
            run = para.add_run(formatted_text)
            run.bold = True
            run.italic = True
            
        elif level == 4:
            # Level 4: Indented, Bold, Title Case, Period.
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
            if not formatted_text.endswith('.'):
                formatted_text += '.'
            run = para.add_run(formatted_text)
            run.bold = True
            run.italic = False
            
        else:  # Level 5
            # Level 5: Indented, Bold Italic, Title Case, Period.
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
            if not formatted_text.endswith('.'):
                formatted_text += '.'
            run = para.add_run(formatted_text)
            run.bold = True
            run.italic = True
        
        # Apply font and ensure black color
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE
        run.font.color.rgb = self.BLACK_COLOR
    
    def _has_first_line_indent(self, para) -> bool:
        """Check if paragraph already has approximately 0.5" first-line indent"""
        indent = para.paragraph_format.first_line_indent
        if indent is None:
            return False
        try:
            # Check if indent is in the tolerance range (0.4" to 0.6")
            return self.INDENT_TOLERANCE_MIN <= indent <= self.INDENT_TOLERANCE_MAX
        except:
            return False
    
    def _format_body_paragraphs(self):
        """Format body paragraphs with APA styling - skip indent if already present"""
        for para in self.doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Skip if this is a heading (already formatted)
            if self._detect_heading_level(para):
                continue
            
            # Skip if this is in a table
            if para._p.getparent().tag.endswith('tc'):
                continue
            
            # Apply body paragraph formatting
            para.paragraph_format.line_spacing = self.LINE_SPACING
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Only indent if not already having an indent AND not a special paragraph
            if not self._has_first_line_indent(para) and not self._is_special_paragraph(para):
                para.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
            
            # Apply font to all runs and ensure black color
            for run in para.runs:
                run.font.name = self.FONT_NAME
                run.font.size = self.FONT_SIZE
                run.font.color.rgb = self.BLACK_COLOR
    
    def _is_special_paragraph(self, para) -> bool:
        """Check if paragraph is special (reference, figure caption, etc.)"""
        text = para.text.strip().lower()
        
        # Reference entries (typically start with author names)
        if para.paragraph_format.left_indent and para.paragraph_format.left_indent > Inches(0.3):
            return True
        
        # Figure/Table captions
        if text.startswith('figure') or text.startswith('table') or text.startswith('note.'):
            return True
        
        return False
    
    def _remove_excessive_blanks(self):
        """Remove excessive blank paragraphs - APA says no extra enters"""
        paragraphs_to_remove = []
        consecutive_blanks = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                consecutive_blanks += 1
                # Allow max 1 blank paragraph between sections
                if consecutive_blanks > 1:
                    paragraphs_to_remove.append(para)
            else:
                consecutive_blanks = 0
        
        # Remove the excessive blank paragraphs
        for para in paragraphs_to_remove:
            try:
                p = para._element
                p.getparent().remove(p)
                self.issues_fixed += 1
            except:
                pass
    
    def _format_references_section(self):
        """Format the references section with proper APA styling"""
        in_references = False
        
        for para in self.doc.paragraphs:
            text = para.text.strip().lower()
            
            # Detect references section header
            if text in ['references', 'reference', 'bibliography', 'works cited']:
                in_references = True
                # Format as section label (centered, bold)
                self._apply_heading_format(para, 1)
                continue
            
            # If we're in references section, apply hanging indent
            if in_references and para.text.strip():
                # Stop if we hit another major heading
                if self._detect_heading_level(para) == 1:
                    in_references = False
                    continue
                
                # Apply hanging indent format
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.left_indent = self.HANGING_INDENT
                para.paragraph_format.line_spacing = self.LINE_SPACING
                
                # Apply hanging indent via XML for reliability
                self._set_hanging_indent(para)
                
                # Apply font
                for run in para.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = self.FONT_SIZE
    
    def _set_hanging_indent(self, para):
        """Set hanging indent using XML for reliability"""
        pPr = para._p.get_or_add_pPr()
        
        # Remove existing ind element if present
        existing_ind = pPr.find(qn('w:ind'))
        if existing_ind is not None:
            pPr.remove(existing_ind)
        
        # Create new ind element with hanging indent
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(self.HANGING_INDENT.twips)))
        ind.set(qn('w:hanging'), str(int(self.HANGING_INDENT.twips)))
        pPr.append(ind)
    
    def _format_tables_and_figures(self):
        """Format tables and figures according to APA guidelines"""
        # Process figures (images with captions)
        figure_num = 0
        table_num = 0
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            # Check for figure captions
            if text.lower().startswith('figure'):
                figure_num += 1
                self._format_figure_caption(para, figure_num)
            
            # Check for table captions  
            elif text.lower().startswith('table'):
                table_num += 1
                self._format_table_caption(para, table_num)
        
        # Format actual tables
        for table in self.doc.tables:
            self._format_table(table)
    
    def _format_figure_caption(self, para, num: int):
        """Format a figure caption according to APA 7th:
        - "Figure X" = Bold only (NOT italic)
        - Title = Italic only (NOT bold)
        """
        text = para.text.strip()
        
        # Extract title from caption (after "Figure X" or "Figure X." or "Figure X:")
        match = re.match(r'figure\s*\d*[\.:]?\s*(.*)', text, re.IGNORECASE)
        title = match.group(1).strip() if match and match.group(1) else ""
        
        # Clear and rebuild
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = self.LINE_SPACING
        
        # "Figure X" in BOLD only (NO italic)
        run1 = para.add_run(f"Figure {num}")
        run1.bold = True
        run1.italic = False  # Explicitly NOT italic
        run1.font.name = self.FONT_NAME
        run1.font.size = self.FONT_SIZE
        run1.font.color.rgb = self.BLACK_COLOR
        
        # Add line break between number and title
        run_break = para.add_run("\n")
        run_break.font.name = self.FONT_NAME
        run_break.font.size = self.FONT_SIZE
        
        # Title in ITALIC only (NO bold)
        if title:
            run2 = para.add_run(self._to_title_case(title))
            run2.italic = True
            run2.bold = False  # Explicitly NOT bold
            run2.font.name = self.FONT_NAME
            run2.font.size = self.FONT_SIZE
            run2.font.color.rgb = self.BLACK_COLOR
    
    def _format_table_caption(self, para, num: int):
        """Format a table caption according to APA 7th:
        - "Table X" = Bold only (NOT italic)
        - Title = Italic only (NOT bold)
        """
        text = para.text.strip()
        
        # Extract title
        match = re.match(r'table\s*\d*[\.:]?\s*(.*)', text, re.IGNORECASE)
        title = match.group(1).strip() if match and match.group(1) else ""
        
        # Clear and rebuild
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.first_line_indent = Inches(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = self.LINE_SPACING
        
        # "Table X" in BOLD only (NO italic)
        run1 = para.add_run(f"Table {num}")
        run1.bold = True
        run1.italic = False  # Explicitly NOT italic
        run1.font.name = self.FONT_NAME
        run1.font.size = self.FONT_SIZE
        run1.font.color.rgb = self.BLACK_COLOR
        
        # Add line break between number and title
        run_break = para.add_run("\n")
        run_break.font.name = self.FONT_NAME
        run_break.font.size = self.FONT_SIZE
        
        # Title in ITALIC only (NO bold)
        if title:
            run2 = para.add_run(self._to_title_case(title))
            run2.italic = True
            run2.bold = False  # Explicitly NOT bold
            run2.font.name = self.FONT_NAME
            run2.font.size = self.FONT_SIZE
            run2.font.color.rgb = self.BLACK_COLOR
    
    def _format_table(self, table):
        """Format a table according to APA guidelines"""
        # APA tables: minimal lines, clean design
        # Apply font to all cells
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = self.FONT_NAME
                        run.font.size = self.FONT_SIZE
    
    def _detect_tables_and_figures(self):
        """Detect all tables and figures in the document for LOT/LOF generation"""
        self.figures_found = []
        self.tables_found = []
        
        figure_num = 0
        table_num = 0
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            text_lower = text.lower()
            
            # Detect figure captions
            if text_lower.startswith('figure'):
                figure_num += 1
                # Extract title
                match = re.match(r'figure\s*\d*\.?\s*(.*)', text, re.IGNORECASE)
                title = match.group(1).strip() if match else text
                if title:
                    self.figures_found.append((figure_num, title))
            
            # Detect table captions
            elif text_lower.startswith('table') and not text_lower.startswith('table of contents'):
                table_num += 1
                # Extract title
                match = re.match(r'table\s*\d*\.?\s*(.*)', text, re.IGNORECASE)
                title = match.group(1).strip() if match else text
                if title:
                    self.tables_found.append((table_num, title))
        
        # Also count actual Word tables if no table captions found
        if not self.tables_found:
            for i, table in enumerate(self.doc.tables, 1):
                # Try to find a title from the first cell or nearby paragraph
                if table.rows and table.rows[0].cells:
                    first_cell_text = table.rows[0].cells[0].text.strip()
                    if first_cell_text:
                        self.tables_found.append((i, f"Table {i}"))
    
    def _generate_front_matter(self):
        """Generate Table of Contents, List of Figures, and List of Tables
        
        Order (per APA for dissertations):
        1. Table of Contents
        2. List of Tables (if 2+ tables)
        3. List of Figures (if 2+ figures)
        
        Each on its own page.
        """
        # Find insertion point (after title page)
        insert_index = self._find_front_matter_insertion_point()
        
        # We'll build front matter in correct order
        # Insert in REVERSE order since each insert pushes content down
        
        # 3. List of Figures (if 2+ figures) - insert last so it appears last
        if len(self.figures_found) >= 2:
            self._insert_list_of_figures(insert_index)
        
        # 2. List of Tables (if 2+ tables)
        if len(self.tables_found) >= 2:
            self._insert_list_of_tables(insert_index)
        
        # 1. Table of Contents - insert first so it appears first
        if self.headings_found:
            self._insert_table_of_contents(insert_index)
    
    def _find_front_matter_insertion_point(self) -> int:
        """Find where to insert front matter (after title page)"""
        # Look for Abstract or first major heading
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip().lower()
            # Find Abstract heading - insert before it
            if text == 'abstract':
                return i
            # Or find first substantial content after title elements
            if i > 5 and text and len(text) > 30:
                return i
        # Default to after first 3 paragraphs (typical title page)
        return min(3, len(self.doc.paragraphs))
    
    def _insert_table_of_contents(self, insert_index: int):
        """Create and insert Table of Contents at specific position"""
        # Filter headings to include only Level 1 and 2
        toc_headings = [(level, text) for level, text, _ in self.headings_found 
                        if level <= 2 and text.lower() not in ['table of contents', 'list of tables', 'list of figures', 'contents']]
        
        if not toc_headings:
            return
        
        # Get body element
        body = self.doc._body._body
        
        # Get reference element for insertion
        if insert_index < len(self.doc.paragraphs):
            ref_element = self.doc.paragraphs[insert_index]._element
        else:
            ref_element = None
        
        # Create TOC title - NOT bold per user request (student papers)
        toc_title = OxmlElement('w:p')
        self._add_paragraph_content(toc_title, "Contents", centered=True, bold=False)
        
        if ref_element is not None:
            ref_element.addprevious(toc_title)
        else:
            body.append(toc_title)
        
        # Add blank line after title
        blank = OxmlElement('w:p')
        toc_title.addnext(blank)
        last_element = blank
        
        # Add each heading entry
        for level, text in toc_headings:
            entry = OxmlElement('w:p')
            clean_text = text.rstrip('.')
            
            # Add indentation for Level 2
            indent = Inches(0.5) if level == 2 else None
            self._add_paragraph_content(entry, clean_text, indent_left=indent)
            
            last_element.addnext(entry)
            last_element = entry
        
        # Add page break after TOC
        page_break = OxmlElement('w:p')
        self._add_page_break(page_break)
        last_element.addnext(page_break)
    
    def _insert_list_of_figures(self, insert_index: int):
        """Create and insert List of Figures"""
        if not self.figures_found:
            return
        
        body = self.doc._body._body
        
        if insert_index < len(self.doc.paragraphs):
            ref_element = self.doc.paragraphs[insert_index]._element
        else:
            ref_element = None
        
        # Create LOF title - NOT bold per user request (student papers)
        lof_title = OxmlElement('w:p')
        self._add_paragraph_content(lof_title, "List of Figures", centered=True, bold=False)
        
        if ref_element is not None:
            ref_element.addprevious(lof_title)
        else:
            body.append(lof_title)
        
        # Add blank line
        blank = OxmlElement('w:p')
        lof_title.addnext(blank)
        last_element = blank
        
        # Add figure entries
        for num, title in self.figures_found:
            entry = OxmlElement('w:p')
            entry_text = f"Figure {num}. {self._to_title_case(title)}"
            self._add_paragraph_content(entry, entry_text)
            last_element.addnext(entry)
            last_element = entry
        
        # Add page break
        page_break = OxmlElement('w:p')
        self._add_page_break(page_break)
        last_element.addnext(page_break)
    
    def _insert_list_of_tables(self, insert_index: int):
        """Create and insert List of Tables"""
        if not self.tables_found:
            return
        
        body = self.doc._body._body
        
        if insert_index < len(self.doc.paragraphs):
            ref_element = self.doc.paragraphs[insert_index]._element
        else:
            ref_element = None
        
        # Create LOT title - NOT bold per user request (student papers)
        lot_title = OxmlElement('w:p')
        self._add_paragraph_content(lot_title, "List of Tables", centered=True, bold=False)
        
        if ref_element is not None:
            ref_element.addprevious(lot_title)
        else:
            body.append(lot_title)
        
        # Add blank line
        blank = OxmlElement('w:p')
        lot_title.addnext(blank)
        last_element = blank
        
        # Add table entries
        for num, title in self.tables_found:
            entry = OxmlElement('w:p')
            entry_text = f"Table {num}. {self._to_title_case(title)}"
            self._add_paragraph_content(entry, entry_text)
            last_element.addnext(entry)
            last_element = entry
        
        # Add page break
        page_break = OxmlElement('w:p')
        self._add_page_break(page_break)
        last_element.addnext(page_break)
    
    def _add_paragraph_content(self, p_element, text: str, centered: bool = False, 
                               bold: bool = False, italic: bool = False,
                               indent_left=None):
        """Add formatted content to a paragraph XML element"""
        # Create paragraph properties
        pPr = OxmlElement('w:pPr')
        
        # Alignment
        if centered:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
        
        # Indentation
        if indent_left:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(int(indent_left.twips)))
            pPr.append(ind)
        
        # Line spacing (double)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '480')  # 480 twips = double spacing
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        
        p_element.append(pPr)
        
        # Create run with text
        run = OxmlElement('w:r')
        
        # Run properties
        rPr = OxmlElement('w:rPr')
        
        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), self.FONT_NAME)
        rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
        rPr.append(rFonts)
        
        # Font size (24 half-points = 12pt)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(szCs)
        
        # Bold
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        
        # Italic
        if italic:
            i = OxmlElement('w:i')
            rPr.append(i)
        
        run.append(rPr)
        
        # Text content
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        run.append(t)
        
        p_element.append(run)
    
    def _add_page_break(self, p_element):
        """Add a page break to a paragraph element"""
        run = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run.append(br)
        p_element.append(run)
    
    def _to_title_case(self, text: str) -> str:
        """Convert text to APA title case"""
        if not text:
            return text
        
        # Words that should not be capitalized (unless first word)
        minor_words = {'a', 'an', 'the', 'and', 'but', 'or', 'nor', 'for', 
                      'yet', 'so', 'at', 'by', 'for', 'in', 'of', 'on', 
                      'to', 'up', 'as', 'is', 'be'}
        
        words = text.split()
        result = []
        
        for i, word in enumerate(words):
            # First word and words after colon are always capitalized
            if i == 0 or (result and result[-1].endswith(':')):
                result.append(word.capitalize())
            elif word.lower() in minor_words:
                result.append(word.lower())
            elif len(word) >= 4:  # Words with 4+ letters are capitalized
                result.append(word.capitalize())
            else:
                result.append(word.lower())
        
        return ' '.join(result)
    
    def save(self, output_path: str):
        """Save the formatted document"""
        self.doc.save(output_path)
    
    def get_issues(self) -> List[Dict]:
        """Get list of formatting issues found"""
        return self.issues_found
    
    def get_issues_fixed_count(self) -> int:
        """Get count of issues fixed"""
        return self.issues_fixed


def format_document(source_path: str, output_path: str, options: Dict) -> Dict:
    """Factory function to format a document"""
    formatter = APAFormatter(source_path)
    formatter.format_document(options)
    formatter.save(output_path)
    
    return {
        'success': True,
        'issues_found': formatter.get_issues(),
        'issues_fixed': formatter.get_issues_fixed_count()
    }
