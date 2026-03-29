"""
DOCX Parser Service
Reads and parses Microsoft Word documents
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Tuple, Optional
import re


class DocxParser:
    """Service for parsing DOCX files"""
    
    def __init__(self, file_path: str):
        self.doc = Document(file_path)
        self.file_path = file_path
    
    def get_all_text(self) -> str:
        """Extract all text from the document"""
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    
    def get_paragraphs(self) -> List[Dict]:
        """Get all paragraphs with their styles and formatting"""
        paragraphs = []
        for i, para in enumerate(self.doc.paragraphs):
            para_info = {
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else None,
                'alignment': str(para.alignment) if para.alignment else None,
                'is_heading': self._is_heading(para),
                'heading_level': self._get_heading_level(para),
                'runs': self._get_runs_info(para)
            }
            paragraphs.append(para_info)
        return paragraphs
    
    def _is_heading(self, para) -> bool:
        """Check if paragraph is a heading"""
        if para.style and 'Heading' in para.style.name:
            return True
        # Check for common heading patterns
        text = para.text.strip()
        if not text:
            return False
        # All caps short text might be heading
        if text.isupper() and len(text) < 100:
            return True
        # Bold short text might be heading
        if len(text) < 100 and para.runs:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold:
                return True
        return False
    
    def _get_heading_level(self, para) -> Optional[int]:
        """Determine the heading level"""
        if para.style:
            style_name = para.style.name
            # Check built-in heading styles
            for i in range(1, 10):
                if f'Heading {i}' in style_name:
                    return min(i, 5)  # APA has 5 levels max
        
        # Infer from formatting
        if self._is_heading(para):
            text = para.text.strip()
            # All caps centered = Level 1
            if text.isupper():
                return 1
            # Bold = Level 2 or 3
            if para.runs and all(run.bold for run in para.runs if run.text.strip()):
                return 2
        return None
    
    def _get_runs_info(self, para) -> List[Dict]:
        """Get formatting info for each run in paragraph"""
        runs = []
        for run in para.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None
            }
            runs.append(run_info)
        return runs
    
    def get_tables(self) -> List[List[List[str]]]:
        """Extract all tables from document"""
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def get_metadata(self) -> Dict:
        """Extract document metadata"""
        core_props = self.doc.core_properties
        return {
            'title': core_props.title,
            'author': core_props.author,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': str(core_props.created) if core_props.created else None,
            'modified': str(core_props.modified) if core_props.modified else None,
        }
    
    def get_word_count(self) -> int:
        """Count words in document"""
        text = self.get_all_text()
        words = re.findall(r'\b\w+\b', text)
        return len(words)
    
    def find_citations(self) -> List[Dict]:
        """Find potential citations in the document"""
        citations = []
        text = self.get_all_text()
        
        # Pattern for parenthetical citations: (Author, Year) or (Author, Year, p. X)
        paren_pattern = r'\([A-Z][a-z]+(?:\s+(?:&|and)\s+[A-Z][a-z]+)*(?:\s+et\s+al\.?)?,?\s*\d{4}[a-z]?(?:,\s*p+\.\s*\d+(?:-\d+)?)?\)'
        
        # Pattern for narrative citations: Author (Year)
        narrative_pattern = r'[A-Z][a-z]+(?:\s+(?:&|and)\s+[A-Z][a-z]+)*(?:\s+et\s+al\.?)?\s*\(\d{4}[a-z]?\)'
        
        for match in re.finditer(paren_pattern, text):
            citations.append({
                'text': match.group(),
                'type': 'parenthetical',
                'start': match.start(),
                'end': match.end()
            })
        
        for match in re.finditer(narrative_pattern, text):
            citations.append({
                'text': match.group(),
                'type': 'narrative',
                'start': match.start(),
                'end': match.end()
            })
        
        return citations
    
    def find_references_section(self) -> Optional[Tuple[int, List[str]]]:
        """Find the references section and extract references"""
        paragraphs = self.get_paragraphs()
        references_start = None
        
        # Find references heading
        for i, para in enumerate(paragraphs):
            text = para['text'].strip().lower()
            if text in ['references', 'reference', 'bibliography', 'works cited']:
                references_start = i
                break
        
        if references_start is None:
            return None
        
        # Collect reference entries
        references = []
        for para in paragraphs[references_start + 1:]:
            text = para['text'].strip()
            if not text:
                continue
            # Stop if we hit another major heading
            if para['is_heading'] and para['heading_level'] == 1:
                break
            references.append(text)
        
        return (references_start, references)


def parse_docx(file_path: str) -> DocxParser:
    """Factory function to create a DocxParser instance"""
    return DocxParser(file_path)
